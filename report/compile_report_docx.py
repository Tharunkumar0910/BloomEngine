import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def build_docx_report(md_path, docx_path, figures_dir):
    print(f"Reading markdown from {md_path}...")
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Standard 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Font Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    lines = content.splitlines()
    in_table = False
    table_lines = []

    def flush_table(t_lines):
        if not t_lines:
            return
        rows_data = []
        for line in t_lines:
            if '|' in line:
                if '---' in line or re.match(r'^\s*\|?\s*:?-+:?\s*\|', line):
                    continue
                parts = [p.strip() for p in line.strip('|').split('|')]
                rows_data.append(parts)

        if not rows_data:
            return

        num_cols = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(rows_data):
            for j, val in enumerate(row):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    cell.text = val
                    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                    
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                    
                    if i == 0:
                        set_cell_background(cell, "1F4E78")
                        for r_item in p.runs:
                            r_item.font.bold = True
                            r_item.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            r_item.font.size = Pt(10)
                    else:
                        bg = "F2F4F8" if i % 2 == 1 else "FFFFFF"
                        set_cell_background(cell, bg)
                        for r_item in p.runs:
                            r_item.font.size = Pt(9.5)
        
        doc.add_paragraph()

    fig_map = {
        "Figure 1.1": "fig_1_1_bloom_taxonomy_hierarchy.png",
        "Figure 1.2": "fig_1_2_system_boundary.png",
        "Figure 4.1": "fig_4_1_system_architecture.png",
        "Figure 4.2": "fig_4_2_system_modules.png",
        "Figure 4.3": "fig_4_3_use_case_diagram.png",
        "Figure 4.4": "fig_4_4_activity_diagram.png",
        "Figure 4.5": "fig_4_5_sequence_diagram.png",
        "Figure 4.6": "fig_4_6_dfd_level1.png",
        "Figure 4.7": "fig_4_7_database_schema.png",
        "Figure 6.1": "fig_6_1_bloom_deberta_distribution.png",
        "Figure 6.2": "fig_6_2_difficulty_deberta_distribution.png",
        "Figure 6.3": "fig_6_3_source_bloom_flan_distribution.png",
        "Figure 6.4": "fig_6_4_target_bloom_flan_distribution.png",
        "Figure 6.5": "fig_6_5_bloom_transformation_matrix.png",
        "Figure 7.1": "fig_7_1_accuracy_f1_epochs.png",
        "Figure 7.2": "fig_7_2_confusion_matrix.png",
        "Figure 7.3": "fig_7_3_normalized_confusion_matrix.png",
        "Figure 7.4": "fig_7_4_roc_curve.png",
        "Figure 7.5": "fig_7_5_loss_curve.png",
    }

    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run("\n".join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if "|" in line and not stripped.startswith("#"):
            in_table = True
            table_lines.append(line)
            continue
        else:
            if in_table:
                flush_table(table_lines)
                table_lines = []
                in_table = False

        if not stripped:
            continue

        if stripped in ["---", "***", "___"]:
            continue

        # Check for Figure Captions to insert high-res figure images
        matched_fig_file = None
        for fig_key, filename in fig_map.items():
            if fig_key in stripped:
                matched_fig_file = os.path.join(figures_dir, filename)
                break

        if matched_fig_file and os.path.exists(matched_fig_file):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(12)
            p_img.paragraph_format.space_after = Pt(4)
            p_img.add_run().add_picture(matched_fig_file, width=Inches(6.0))

        # Headings
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[3:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[4:].strip())
            run.font.name = 'Arial'
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        elif stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            parts = re.split(r'(\*\*.*?\*\*)', bullet_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    p.add_run(part)

        elif re.match(r'^\d+\.\s', stripped):
            num_text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            
            parts = re.split(r'(\*\*.*?\*\*)', num_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    p.add_run(part)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15

            if matched_fig_file:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                    r = p.add_run(part[1:-1])
                    r.font.italic = True
                    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                else:
                    p.add_run(part)

    if in_table:
        flush_table(table_lines)

    doc.save(docx_path)
    print(f"Successfully generated DOCX report with embedded figures at: {docx_path}")

if __name__ == "__main__":
    md_file = r"c:\Tharun\BloomAI_Arena_v2_1\report\BloomEngine_MCA_Major_Project_Report.md"
    docx_file = r"c:\Tharun\BloomAI_Arena_v2_1\report\BloomEngine_MCA_Major_Project_Report.docx"
    fig_dir = r"c:\Tharun\BloomAI_Arena_v2_1\report\figures"
    build_docx_report(md_file, docx_file, fig_dir)
